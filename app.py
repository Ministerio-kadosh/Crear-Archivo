from flask import Flask, render_template, request, make_response
from docx import Document
import io
import os
from datetime import date
from email import enviar_por_correo  # Asumo que existe

app = Flask(__name__)

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/generar', methods=['POST'])
def generar():
    data = request.get_json()
    nombre_usuario = data.get('nombre', 'SinNombre')

    fecha = date.today().isoformat()
    nombre_archivo = f"informe_{nombre_usuario}_{fecha}.docx"

    # Crear documento Word
    doc = Document()
    doc.add_heading('Informe de Usuario', 0)
    doc.add_paragraph(f'Nombre: {nombre_usuario}')

    table = doc.add_table(rows=2, cols=3)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Producto'
    hdr_cells[1].text = 'Cantidad'
    hdr_cells[2].text = 'Precio'

    row_cells = table.rows[1].cells
    row_cells[0].text = 'Chetos'
    row_cells[1].text = '2'
    row_cells[2].text = '$20'

    # Guardar en memoria
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Enviar por correo (ojo que enviar_por_correo debe aceptar bytes o file-like)
    correos = ['adalig930@gmail.com', 'correo2@gmail.com']
    enviar_por_correo(buffer, correos)
    buffer.seek(0)  # reset para la descarga

    response = make_response(buffer.read())
    response.headers.set('Content-Disposition', 'attachment', filename=nombre_archivo)
    response.headers.set('Content-Type', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    return response

if __name__ == '__main__':
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port)
